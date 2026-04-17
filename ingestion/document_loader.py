"""
Multi-format document loader.

Supports: PDF (digital + scanned), DOCX, XLSX, plain text.
Output: a list of page-level text chunks ready for the extraction pipeline.

Dependencies (install via requirements.txt):
  pdfplumber, python-docx, openpyxl, Pillow
OCR fallback uses pytesseract (requires Tesseract binary).
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class DocumentPage:
    page_number: int
    text: str
    is_ocr: bool = False


@dataclass
class LoadedDocument:
    filename: str
    pages: List[DocumentPage]
    raw_format: str                        # "pdf" | "docx" | "xlsx" | "txt"
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


class DocumentLoader:
    """
    Loads a collateral schedule document from disk and returns a LoadedDocument.

    Usage:
        loader = DocumentLoader(ocr_fallback=True)
        doc = loader.load("path/to/csa.pdf")
    """

    # Minimum character count on a PDF page before we consider it a scan.
    _OCR_THRESHOLD = 50

    def __init__(self, ocr_fallback: bool = True):
        self.ocr_fallback = ocr_fallback

    def load(self, path: str | Path) -> LoadedDocument:
        path = Path(path)
        suffix = path.suffix.lower()

        dispatch = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".doc": self._load_docx,
            ".xlsx": self._load_xlsx,
            ".xls": self._load_xlsx,
            ".txt": self._load_txt,
        }

        loader_fn = dispatch.get(suffix)
        if loader_fn is None:
            raise ValueError(f"Unsupported file type: {suffix}")

        logger.info("Loading %s (%s)", path.name, suffix)
        return loader_fn(path)

    # ── Format-specific loaders ───────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> LoadedDocument:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pip install pdfplumber")

        pages: List[DocumentPage] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                is_ocr = False

                if len(text.strip()) < self._OCR_THRESHOLD and self.ocr_fallback:
                    logger.debug("Page %d appears scanned — falling back to OCR", i)
                    img = page.to_image(resolution=300).original
                    from ingestion.ocr_processor import OCRProcessor
                    text = OCRProcessor().image_to_text(img)
                    is_ocr = True

                pages.append(DocumentPage(page_number=i, text=text, is_ocr=is_ocr))

        return LoadedDocument(
            filename=path.name,
            pages=pages,
            raw_format="pdf",
            metadata={"path": str(path)},
        )

    def _load_docx(self, path: Path) -> LoadedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx")

        doc = Document(str(path))
        # DOCX has no concept of pages in the DOM; treat as single page.
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        return LoadedDocument(
            filename=path.name,
            pages=[DocumentPage(page_number=1, text=text)],
            raw_format="docx",
        )

    def _load_xlsx(self, path: Path) -> LoadedDocument:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("pip install openpyxl")

        wb = openpyxl.load_workbook(str(path), data_only=True)
        pages: List[DocumentPage] = []

        for sheet_idx, sheet in enumerate(wb.worksheets, start=1):
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
            text = "\n".join(rows)
            pages.append(DocumentPage(page_number=sheet_idx, text=text))

        return LoadedDocument(
            filename=path.name,
            pages=pages,
            raw_format="xlsx",
        )

    def _load_txt(self, path: Path) -> LoadedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(
            filename=path.name,
            pages=[DocumentPage(page_number=1, text=text)],
            raw_format="txt",
        )


def chunk_document(doc: LoadedDocument, max_tokens: int = 3000) -> List[str]:
    """
    Split a LoadedDocument into overlapping chunks suitable for LLM context windows.

    Simple word-count heuristic: 1 token ≈ 0.75 words.
    Real implementation should use the tokenizer of the target model.
    """
    max_words = int(max_tokens * 0.75)
    overlap_words = max_words // 5

    words = doc.full_text.split()
    chunks: List[str] = []
    start = 0

    while start < len(words):
        end = min(start + max_words, len(words))
        chunks.append(" ".join(words[start:end]))
        start += max_words - overlap_words

    return chunks
